import argparse
import docx
from pathlib import Path
import win32com.client
from tqdm import tqdm
import helpers as hp
import time

def replaceTags(document, filepath):
    for i in range(len(document.paragraphs)):
        line_str = document.paragraphs[i].text
        if line_str.endswith('<br>'):
            continue
        line_str += '<br>'
        document.paragraphs[i].text = line_str
    document.save(str(filepath))

def getTags(paths):
    word = win32com.client.Dispatch("Word.Application")
    # wdFormatDocumentDefault = 16

    if paths['bulk']:
        # Convert all .doc files to .docx
        for filepath in tqdm(sorted(Path(paths['input']).glob("*.doc"))):
            # if str(filepath).endswith(".doc"):
            hp.saveAsDocx(word, filepath)
        
        for filepath in tqdm(sorted(Path(paths['input']).glob("*.docx"))):
            document = docx.Document(str(filepath))

            replaceTags(document, filepath)

    else:
        filepath = Path(paths['input'])
        if str(filepath).endswith(".doc") or str(filepath).endswith(".DOC"):
            hp.saveAsDocx(word, filepath)
        
        document = docx.Document(str(filepath.parent / filepath.stem) + ".docx")

        replaceTags(document, filepath)


def main():
    # Initialise parser
    parser = argparse.ArgumentParser(description="Tool to replace new line characters with <br> tags in word file")

    # Add arguments
    parser.add_argument('inpath', help="Path of the file or folder.")
    # parser.add_argument('--rm', default=False, help="Set to True, if in-place operation is needed")

    args = parser.parse_args()

    # print(args.rm)
    start = time.time()
    paths = hp.resolvePath(args.inpath, None)
    getTags(paths)
    end = time.time()
    print("Time taken: ", end-start)

main()